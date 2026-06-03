"""Auto-generate a resume from live DB data, matching the real Word doc layout."""

import io
import subprocess
import tempfile
from pathlib import Path

from fastapi import APIRouter, Depends
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.database import get_db
from app import models

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

router = APIRouter(prefix="/resume", tags=["resume"])

FONT_NAME = "Calibri"
FONT_SIZE = Pt(10.5)
NAME_SIZE = Pt(18)
TAB_POS = Inches(7.50)


def _set_bottom_border(paragraph):
    """Add a bottom border to a paragraph (like the section headers)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {qn("w:xmlns")}="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_right_tab(paragraph):
    """Add a right-aligned tab stop at 7.5 inches."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_POS, WD_TAB_ALIGNMENT.RIGHT)


def _run(paragraph, text, bold=False, italic=False, size=None, color=None, name=None):
    """Add a run with specific formatting."""
    r = paragraph.add_run(text)
    r.font.name = name or FONT_NAME
    r.font.size = size or FONT_SIZE
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    return r


def _para(doc, text="", bold=False, alignment=None, spacing_after=Pt(0), spacing_before=Pt(0)):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = spacing_before
    p.paragraph_format.line_spacing = Pt(13)
    if text:
        _run(p, text, bold=bold)
    return p


async def _build_docx(db: AsyncSession) -> io.BytesIO:
    about_r = await db.execute(select(models.About).where(models.About.id == 1))
    about = about_r.scalar_one_or_none()

    exp_r = await db.execute(select(models.Experience).order_by(models.Experience.sort_order))
    experiences = exp_r.scalars().all()

    skills_r = await db.execute(select(models.Skill).order_by(models.Skill.sort_order))
    skills = skills_r.scalars().all()

    projects_r = await db.execute(select(models.Project).order_by(models.Project.sort_order))
    projects = projects_r.scalars().all()

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.5)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Name ────────────────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_bottom_border(p)
    _run(p, "Nathan Blatter", bold=True, size=NAME_SIZE)

    # ── Contact ─────────────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(2))
    _set_bottom_border(p)
    _run(p, "(925) 886-9553 | nzb22@byu.edu | linkedin.com/in/nathanblatter")

    # ── Summary ─────────────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=Pt(4))
    _set_bottom_border(p)
    _run(p, "Information Systems student (Full-Stack Software Engineering emphasis)", bold=True)
    _run(p, " with experience in C#, Java, Python, SQL, and cloud technologies. "
            "Proven ability to build full-stack applications, AI-driven systems, and data pipelines. "
            "Known for strong ownership, clean code practices, and rapid skill acquisition.")

    # ── Education ───────────────────────────────────────────────────────
    edu_entries = [e for e in experiences if any(k in e.title for k in ("B.S.", "M.S.", "Bachelor", "Master"))]
    work_entries = [e for e in experiences if e not in edu_entries]

    p = _para(doc, spacing_before=Pt(4))
    _set_bottom_border(p)
    _run(p, "EDUCATION", bold=True)

    for edu in edu_entries:
        # Title with right-aligned date
        p = _para(doc)
        _add_right_tab(p)
        date_part = edu.year.split("—")[-1].strip() if "—" in edu.year else edu.year
        _run(p, f"{edu.title}\t{date_part}", bold=True)

        # Subtitle and description lines
        for line in edu.description.split(". "):
            line = line.strip().rstrip(".")
            if line:
                _para(doc, line)

    # ── Technical Skills ────────────────────────────────────────────────
    p = _para(doc, spacing_before=Pt(4))
    _set_bottom_border(p)
    _run(p, "TECHNICAL SKILLS", bold=True)

    cats: dict[str, list[str]] = {}
    for sk in skills:
        cats.setdefault(sk.category, []).append(sk.name)
    for cat, names in cats.items():
        p = _para(doc)
        _run(p, f"{cat}: ", bold=True)
        _run(p, ", ".join(names))

    # ── Projects ────────────────────────────────────────────────────────
    p = _para(doc, spacing_before=Pt(4))
    _set_bottom_border(p)
    _run(p, "PROJECTS", bold=True)

    for proj in projects[:6]:
        tags_str = f" ({', '.join(proj.tags[:4])})" if proj.tags else ""
        year = proj.year or ""
        p = _para(doc)
        _run(p, f"{proj.title}{tags_str} | {year}", bold=True)

        desc_lines = [s.strip() for s in proj.description.split(". ") if s.strip()]
        for line in desc_lines[:3]:
            line = line.rstrip(".")
            _para(doc, f"• {line}")

    # ── Experience ──────────────────────────────────────────────────────
    p = _para(doc, spacing_before=Pt(4))
    _set_bottom_border(p)
    _run(p, "EXPERIENCE", bold=True)

    for exp in work_entries:
        # Title with right-aligned date
        p = _para(doc)
        _add_right_tab(p)
        _run(p, f"{exp.title}\t{exp.year}", bold=True)

        # Subtitle (org name)
        _para(doc, exp.subtitle)

        # Description bullets
        desc_lines = [s.strip() for s in exp.description.split(". ") if s.strip()]
        for line in desc_lines[:4]:
            line = line.rstrip(".")
            _para(doc, f"• {line}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/pdf")
async def resume_pdf(db: AsyncSession = Depends(get_db)):
    docx_buf = await _build_docx(db)

    # Try converting to PDF via LibreOffice if available
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "resume.docx"
            docx_path.write_bytes(docx_buf.read())
            docx_buf.seek(0)

            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, str(docx_path)],
                capture_output=True, timeout=30,
            )
            pdf_path = Path(tmpdir) / "resume.pdf"
            if result.returncode == 0 and pdf_path.exists():
                return Response(
                    content=pdf_path.read_bytes(),
                    media_type="application/pdf",
                    headers={"Content-Disposition": "inline; filename=nathan-blatter-resume.pdf"},
                )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: serve docx directly
    docx_buf.seek(0)
    return StreamingResponse(
        docx_buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "inline; filename=nathan-blatter-resume.docx"},
    )


@router.get("/docx")
async def resume_docx(db: AsyncSession = Depends(get_db)):
    docx_buf = await _build_docx(db)
    return StreamingResponse(
        docx_buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=nathan-blatter-resume.docx"},
    )
